from docx import Document
import os

def replace_dashes_with_asterisks(docx_file):
    document = Document(docx_file)
    for paragraph in document.paragraphs:
        # Split the paragraph text by newline characters
        lines = paragraph.text.split('\n')
        for i, line in enumerate(lines):
            # Search for dashes from 20 down to 3
            for dash_count in range(20, 2, -1):
                # Construct the dash sequence
                dash_sequence = '-' * dash_count
                # Check if the dash sequence is present in the line
                if dash_sequence in line:
                    # Replace the dash sequence with three asterisks
                    line = line.replace(dash_sequence, '***')
                    # Update the line in the list of lines
                    lines[i] = line
                    # Break the loop as we've found and replaced the sequence
                    break
        # Join the modified lines back into a paragraph
        modified_paragraph_text = '\n'.join(lines)
        # Clear the original paragraph and add the modified text
        paragraph.clear()
        paragraph.add_run(modified_paragraph_text)
    
    # Save the modified document
    modified_file = docx_file.replace('.docx', '_modified.docx')
    document.save(modified_file)
    print(f"Modified document saved as: {modified_file}")

def process_all_docx_in_folder(folder_path):
    for filename in os.listdir(folder_path):
        if filename.endswith('.docx') and not filename.endswith('_modified.docx'):
            file_path = os.path.join(folder_path, filename)
            print(f"Processing file: {file_path}")
            replace_dashes_with_asterisks(file_path)

# Get the folder path of the current script
folder_path = os.path.dirname(os.path.abspath(__file__))

# Process all .docx files in the folder
process_all_docx_in_folder(folder_path)
